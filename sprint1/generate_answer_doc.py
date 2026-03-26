"""Generate concise, API-focused answer.docx for CSE 416 Data Design Activity"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

NAVY  = RGBColor(0x00, 0x33, 0x99)
GREY  = RGBColor(0x44, 0x44, 0x44)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs: r.font.color.rgb = NAVY
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs: r.font.color.rgb = NAVY
    return p

def para(text, bold=False, italic=False, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.size = Pt(size)
    return p

def bullet(label, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    r.bold = True; r.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x6e)

def small_table(headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Light Shading Accent 1'
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.size = Pt(9)
    for row in rows:
        r = tbl.add_row()
        for i, val in enumerate(row):
            r.cells[i].text = val
            r.cells[i].paragraphs[0].runs[0].font.size = Pt(9)

# ── Title ─────────────────────────────────────────────────────────────────────
t = doc.add_heading('Dark Blue Draft Kit — Data Design', 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in t.runs: r.font.color.rgb = NAVY

sub = doc.add_paragraph('CSE 416 | Team Activity: Data Design | 2026-03-26')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.color.rgb = GREY
sub.runs[0].font.size = Pt(10)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
h1('1. Data Sources')
# ═══════════════════════════════════════════════════════════════════════════════
para('All data sourced from the MLB Stats API (statsapi.mlb.com). Static data is cached in players.json; '
     'live data fetched via nightly/hourly cron jobs during the MLB season (April–October).', size=10)
doc.add_paragraph()

small_table(
    ['Data', 'MLB Stats API Endpoint', 'Key Fields', 'Refresh'],
    [
        ('Player info',
         '/v1/sports/1/players?season=2025&gameType=R',
         'id (MLBAM), fullName, currentTeam.id, primaryPosition',
         'Once per season'),
        ('2025 stats (hitting)',
         '/v1/people/{id}/stats?stats=season&season=2025&group=hitting',
         'homeRuns, rbi, stolenBases, avg, obp, slg, runs',
         'Daily (in-season)'),
        ('2025 stats (pitching)',
         '/v1/people/{id}/stats?stats=season&season=2025&group=pitching',
         'era, whip, strikeOuts, wins, saves',
         'Daily (in-season)'),
        ('Previous seasons',
         '/v1/people/{id}/stats?stats=yearByYear&group=hitting,pitching',
         'Same fields, keyed by season year',
         'Static after season ends'),
        ('Injury / roster status',
         '/v1/transactions?sportId=1&startDate={today}&endDate={today}',
         'status (Active, IL-10, IL-60), injuryDescription',
         'Every 4 hours'),
        ('Team info',
         '/v1/teams?sportId=1',
         'id (MLBAM), name, abbreviation, league.id',
         'Once per season'),
        ('Depth charts',
         '/v1/teams/{team_id}/roster?rosterType=depthChart&season=2025',
         'roster[].player.id, position, status',
         'Daily'),
        ('Transactions',
         '/v1/transactions?sportId=1&startDate={N_days_ago}&endDate={today}',
         'type (OPTION, RECALL, DFA, TRADE, SIGN, RELEASE)',
         'Every 4 hours'),
        ('Player photos',
         'img.mlbstatic.com/mlb-photos/image/.../people/{id}/headshot/67/current',
         'Constructed from MLBAM id — no separate call',
         'On-demand (CDN)'),
    ]
)
doc.add_paragraph()
para('Players.json is rebuilt by a Node.js script (scripts/generate-players.js) that calls the MLB Stats API '
     'and merges all fields. 313 players are currently in the dataset.', italic=True, size=9.5)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
h1('2. Player and Team ID Management')
# ═══════════════════════════════════════════════════════════════════════════════
para('Player names are NOT unique keys — two players may share a name (including on the same team). '
     'All objects are keyed by MLBAM integer ID, the same ID used by the official MLB Stats API, '
     'Baseball Reference, and FanGraphs. This ID is stable across teams, seasons, and name changes.', size=10)
doc.add_paragraph()

small_table(
    ['Object', 'Primary Key', 'Display Field', 'Example'],
    [
        ('Player', 'MLBAM integer id  (e.g. 660271)', 'fullName', '660271 = Shohei Ohtani'),
        ('Team',   'MLBAM team id     (e.g. 119)',    'abbreviation', '119 = LAD'),
        ('League', 'MLBAM league id   (103 AL / 104 NL)', 'name', '103 = American League'),
    ]
)
doc.add_paragraph()
para('Player photos use the MLBAM ID directly in the CDN URL — no separate lookup needed. '
     'The current players.json uses sequential integers (1, 2, 3…) as internal IDs since the '
     'dataset is a projected/synthesized 2025 roster; production will replace these with official MLBAM IDs.', italic=True, size=9.5)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
h1('3. Player Valuation API — Request Signatures')
# ═══════════════════════════════════════════════════════════════════════════════
para('Base URL: https://draftapi.anythingavenue.com  |  Auth: X-License-Key header  |  Demo key: DB-2026-DEMO-0001', bold=True, size=10)
doc.add_paragraph()

# 3.1
h2('3.1  $ Value of a Single Player — POST /v1/valuate')
para('Returns the draft-state-aware auction dollar value for one nominated player.', size=10)
code('curl -X POST https://draftapi.anythingavenue.com/v1/valuate \\\n'
     '  -H "X-License-Key: DB-2026-DEMO-0001" \\\n'
     '  -H "Content-Type: application/json" \\\n'
     '  -d \'{"draft_state":{"total_teams":12,"budget_per_team":260,\n'
     '         "scoring_categories":["HR","RBI","AVG","SB","ERA","SO","WHIP"],\n'
     '         "nominated_player":"Shohei Ohtani","teams":[],\n'
     '         "roster_config":{"C":1,"1B":1,"2B":1,"3B":1,"SS":1,"OF":3,"SP":2,"RP":2,"UTIL":1,"BN":2}}}\'')
para('Response (200 OK) — confirmed live:', size=10)
code('{"player":"Shohei Ohtani","true_dollar_value":75,"max_bid_recommendation":69,\n'
     ' "market_inflation":1.0,"scarcity_tier":"LOW",\n'
     ' "position_scarcity":{"DH":"LOW","SP":"LOW"},"draftability_score":1,\n'
     ' "reasoning":"Tier: LOW. TDV: $75.",\n'
     ' "stats":{"tier":"Elite","positions":["DH","SP"],"team":"LAD","league":"NL"}}')

doc.add_paragraph()
small_table(
    ['Field', 'Type', 'Description'],
    [
        ('total_teams', 'int', 'Number of teams in the league (affects budget inflation)'),
        ('budget_per_team', 'int', 'Auction budget per team (typically 260)'),
        ('scoring_categories', 'string[]', 'Active scoring cats: HR, RBI, AVG, SB, ERA, SO, WHIP, etc.'),
        ('nominated_player', 'string', 'Player name to valuate (case-insensitive, partial match supported)'),
        ('teams', 'object[]', 'Array of {id, budget_remaining, roster:[names...]} — current draft state'),
        ('roster_config', 'object', 'Slots per position: {C:1, 1B:1, 2B:1, 3B:1, SS:1, OF:3, SP:2, RP:2, UTIL:1, BN:2}'),
    ]
)
doc.add_paragraph()

# 3.2
h2('3.2  $ Value of a Player Collection — POST /v1/valuate/batch')
para('Valuates up to 50 players in a single request. Same draft_state as above, plus a players array.', size=10)
code('curl -X POST https://draftapi.anythingavenue.com/v1/valuate/batch \\\n'
     '  -H "X-License-Key: DB-2026-DEMO-0001" -H "Content-Type: application/json" \\\n'
     '  -d \'{"players":["Shohei Ohtani","Juan Soto","Kyle Tucker"],\n'
     '         "draft_state":{"total_teams":12,"budget_per_team":260,\n'
     '           "scoring_categories":["HR","RBI","AVG","SB","ERA","SO","WHIP"],\n'
     '           "teams":[],"roster_config":{"C":1,"1B":1,"2B":1,"3B":1,"SS":1,"OF":3,"SP":2,"RP":2,"UTIL":1,"BN":2}}}\'')
para('Response (200 OK):', size=10)
code('{"count":3,\n'
     ' "valuations":[\n'
     '   {"player":"Shohei Ohtani","true_dollar_value":75,"max_bid_recommendation":69,...},\n'
     '   {"player":"Juan Soto","true_dollar_value":66,"max_bid_recommendation":61,...},\n'
     '   {"player":"Kyle Tucker","true_dollar_value":57,"max_bid_recommendation":52,...}\n'
     ' ],\n'
     ' "errors":[]}')
doc.add_paragraph()

# 3.3
h2('3.3  $ Value of All Eligible Players — GET /v1/valuate/all')
para('Returns all undrafted players valued and sorted by true_dollar_value descending. '
     'Pass drafted player names via the drafted param to exclude them from the pool.', size=10)
code('curl "https://draftapi.anythingavenue.com/v1/valuate/all\n'
     '  ?total_teams=12&budget_per_team=260\n'
     '  &scoring_categories=HR,RBI,AVG,SB,ERA,SO,WHIP\n'
     '  &drafted=Shohei+Ohtani,Juan+Soto\n'
     '  &pos=OF&tier=Elite" \\\n'
     '  -H "X-License-Key: DB-2026-DEMO-0001"')
para('Response (200 OK):', size=10)
code('{"count":311,\n'
     ' "draft_state_summary":{"total_teams":12,"budget_per_team":260,"drafted_count":2},\n'
     ' "players":[\n'
     '   {"player":"Kyle Tucker","true_dollar_value":57,"max_bid_recommendation":52,...},\n'
     '   {"player":"Trea Turner","true_dollar_value":53,"max_bid_recommendation":48,...},\n'
     '   ...311 players sorted by true_dollar_value desc\n'
     ' ]}')
doc.add_paragraph()
para('All three endpoints return context-aware values — true_dollar_value changes as teams draft players, '
     'budgets deplete, and positional scarcity shifts.', italic=True, size=9.5)

# ── Save ──────────────────────────────────────────────────────────────────────
out = r'C:\Users\Apple\Desktop\answer.docx'
doc.save(out)
print(f'Saved: {out}')
